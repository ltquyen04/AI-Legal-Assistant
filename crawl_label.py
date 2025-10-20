import os
import re
import pandas as pd
from docx import Document
from tqdm import tqdm

def extract_info_from_table(doc):
    """Trích xuất luật số và thời gian ban hành từ bảng đầu (nếu có)."""
    so_luat = ''
    thoi_gian = ''
    if doc.tables:
        table = doc.tables[0]
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()

                # Tìm Số hoặc Luật số
                if re.search(r'\b(Số|Luật số)\b', text, re.IGNORECASE):
                    so_luat = text.split(':', 1)[1].strip()                      

                # Tìm ngày ban hành
                if re.search(r'ngày\s+\d{1,2}\s+tháng\s+\d{1,2}\s+năm\s+\d{4}', text, re.IGNORECASE):
                    match = re.search(r'ngày\s+\d{1,2}\s+tháng\s+\d{1,2}\s+năm\s+\d{4}', text, re.IGNORECASE)
                    if match:
                        thoi_gian = match.group(0)

    return so_luat, thoi_gian

def extract_loai_van_ban(so_text):
    if 'NĐ' in so_text:
        return 'Nghị định'
    elif 'TT' in so_text:
        return 'Thông tư'
    return 'Luật'

def extract_dieu_luat(texts):
    dieu_pattern = re.compile(r'^Điều\s+(\d+)[.:]?\s*(.*)', re.IGNORECASE)
    data = []
    current_dieu = None

    for line in texts:
        line = line.strip()
        if not line:
            continue

        match = dieu_pattern.match(line)
        if match:
            if current_dieu:
                data.append(current_dieu)
            current_dieu = {
                'dieu': match.group(1),
                'ten_dieu': match.group(2).strip() if match.group(2) else '',
                'noi_dung': ''
            }
        elif current_dieu:
            current_dieu['noi_dung'] += line + ' '

    if current_dieu:
        data.append(current_dieu)
    return data

def process_docx(filepath):
    doc = Document(filepath)
    all_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    so, thoi_gian = extract_info_from_table(doc)
    loai_van_ban = extract_loai_van_ban(so)
    dieu_luat = extract_dieu_luat(all_texts)

    rows = []
    for dieu_info in dieu_luat:
        row = {
            'id': f"{so}_{dieu_info['dieu']}",
            'Loại tài liệu pháp luật': loai_van_ban,
            'số': so,
            'điều': dieu_info['dieu'],
            'tên điều': dieu_info['ten_dieu'],
            'nội dung': dieu_info['noi_dung'].strip(),
            'thời gian ban hành': thoi_gian,
            'nhãn': ''
        }
        rows.append(row)
    return rows

def crawl_folder_to_csv(folder_path='data_retriever', output_csv_path='data/legal_data.csv'):
    all_rows = []

    for filename in tqdm(os.listdir(folder_path), desc="Đang xử lý"):
        if filename.endswith('.docx'):
            filepath = os.path.join(folder_path, filename)
            try:
                rows = process_docx(filepath)
                all_rows.extend(rows)
            except Exception as e:
                print(f"❌ Lỗi xử lý {filename}: {e}")

    df = pd.DataFrame(all_rows)
    df.to_csv(output_csv_path, index=False, encoding='utf-8-sig')
    print(f"✅ Đã lưu CSV vào: {output_csv_path}")

def main():
    folder_input = 'data_retriever'
    output_file = 'data/legal_data.csv'
    crawl_folder_to_csv(folder_input, output_file)

if __name__ == '__main__':
    main()
